import os
import pandas as pd
import docx
from docx import Document
import re
import xlsxwriter
# from xlwt  import Workbook

##---- CONFIGURATION VARIABLES ----##
Outputfilename="Extraction_outputV1.xlsx"
#INIT
init_word_1= 'Headline'
stop_word_1= 'What'

#TWO
start_word2='What'
start_word_2a='What'
start_word_2b='What IFC'
start_word_2c='xxx'
start_word_2d= 'xxx'

stop_word_2a= 'Based'
stop_word_2b= 'xxx'
stop_word_2c= 'What actually'
stop_word_2d = init_word_1
stop_word_2e= 'xxx'

#THREE
start_word_3='Based'
start_word_3a= 'What actually' 
start_word_3b='xxx'
start_word_3c='xxx'
start_word_3d='xxx'
start_word_3e='xxx'

stop_word_3a='Lesson'
stop_word_3b='what advice'
stop_word_3c= init_word_1
stop_word_3d= 'xxx'
stop_word_3e='Annex'

#FOUR
start_word_4='Lesson'
start_word_4a= 'what advice'
start_word_4b='lesson for'
start_word_4c='xxx'
start_word_4d='xxx'

stop_word_4a= init_word_1
stop_word_4b= 'Annex'
stop_word_4c= 'Area/Theme'
stop_word_4d= "related"

##----- CONFIGURATION END -----##

strPath = os.getcwd()
# strPath = 'Z:\\1. Micro Evaluation\\VELOCITY\\paragraph_predict\\XPSR\\test\\test2\\wrong'
# doc = docx.AdvSearch()

translist = list()

### --- Directory Crawl --- ###
for file in os.listdir(strPath):
	if file.endswith('.docx'):
		doc_id = re.split('_|\s', os.path.splitext(file)[0])[0]
		print("_________\ncreated document_id: {} for: \n{}".format(doc_id,file))
		doc = docx.Document(file)
		sublesscounter=0
		# paralist = []
		for i in range(len(doc.paragraphs)):
			headline=''
			lesson_id=None
			what=''
			based=''
			lesson=''
			if re.match(r'^%s\b'% init_word_1,doc.paragraphs[i].text): #.startswith('Headline') == True:
				sublesscounter+= 1
				lesson_id = str(doc_id+'-'+str(sublesscounter))
				
				print('_________\ninit.DEBUG:----> ',i)
				e=i
				# o=i
				# u=i

				while doc.paragraphs[e].text.startswith(stop_word_1)==False:
					headline = headline+' - '+doc.paragraphs[e].text
					print('H.DEBUG:----> ',e)
					e+=1
				if doc.paragraphs[e].text.startswith(start_word_2a)==True or doc.paragraphs[e].text.startswith(start_word_2b)==True or doc.paragraphs[e].text.startswith(start_word_2c)==True or doc.paragraphs[e].text.startswith(start_word_2d)==True:
					while doc.paragraphs[e].text.startswith(stop_word_2a)==False and doc.paragraphs[e].text.startswith(stop_word_2b)==False and doc.paragraphs[e].text.startswith(stop_word_2c)==False and doc.paragraphs[e].text.startswith(stop_word_2d)==False and doc.paragraphs[e].text.startswith(stop_word_2e)==False:
						what = what+' - ' +doc.paragraphs[e].text
						print('W.DEBUG:----> ',e)
						e+=1				
				if doc.paragraphs[e].text.startswith(start_word_3a)==True or doc.paragraphs[e].text.startswith(start_word_3b)==True or doc.paragraphs[e].text.startswith(start_word_3c)==True or doc.paragraphs[e].text.startswith(start_word_3d)==True or doc.paragraphs[e].text.startswith(start_word_3e)==True:
					while doc.paragraphs[e].text.startswith(stop_word_3a)==False and doc.paragraphs[e].text.startswith(stop_word_3b)==False and doc.paragraphs[e].text.startswith(stop_word_3c)==False and doc.paragraphs[e].text.startswith(stop_word_3d)==False and doc.paragraphs[e].text.startswith(stop_word_3e)==False:
						print('B.DEBUG:----> ',e)
						based = based+'-'+doc.paragraphs[e].text
						e+=1
				if doc.paragraphs[e].text.startswith(start_word_4a)==True or doc.paragraphs[e].text.startswith(start_word_4b)==True and doc.paragraphs[e].text.startswith(start_word_4c)==True and doc.paragraphs[e].text.startswith(start_word_4d)==True:
					while doc.paragraphs[e].text.startswith(stop_word_4a)==False and doc.paragraphs[e].text.startswith(stop_word_4b)==False and doc.paragraphs[e].text.startswith(stop_word_4c)==False and doc.paragraphs[e].text.startswith(stop_word_4d)==False:
						print('L.DEBUG:----> ',e)
						lesson = lesson+'-'+doc.paragraphs[e].text
						e+=1
					
					# paradex = doc.paragraph[i].text.find('What')
					# tillwhat = doc.paragraphs[i:].find('What')

					# while i<tillwhat:
					# 	headline = headline+'-'+doc.paragraphs[i].text
					# 	i+=1
					# what = doc.paragraphs[tillwhat]					 


				print("--\n>>{} captured as: {}\n--\n>>{} captured as: {}\n--\n>>{} captured as: {}\n--\n>>{} captured as: {}\n".format(init_word_1.upper(),headline,start_word_2a.upper(),what,start_word_3a.upper(),based,start_word_4a.upper(),lesson))
				sub_row_list = [doc_id,lesson_id,headline,what,based,lesson] #doc.paragraphs[i].text]
				translist.append(sub_row_list)
			else:
				pass
			
			# print("doc_id: {} \nlesson_id: {}\n paragraphs: \n".format(doc_id,lesson_id),doc.paragraphs[i].text)

		print("created lessons for {}\n\n--".format(file))

	else:
		pass
print("||Directory Crawl Complete")


#FROM LIST TO DATAFRAME
df=pd.DataFrame(translist, columns=['id','lesson_id','headline','what','based','lesson'])

####----EXCEL WRITER----####

writer = pd.ExcelWriter(strPath+'\\'+Outputfilename)
df.to_excel(writer, sheet_name='Sheet1', startrow=0, startcol=0)
# workbook = xlsxwriter.Workbook('Testoutput.xlsx')
# workbook.autofilter('A1:e1')
writer.save()

print("||output document saved as '{} in {}'".format(Outputfilename,strPath))
print('||test complete')